from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_report():
    doc = Document()

    # Configuración de estilos para que parezca profesional
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # 1. PORTADA (Página 1)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n' * 3)
    run = title.add_run('TRABAJO DE ENFOQUE')
    run.bold = True
    run.size = Pt(26)
    
    mod = doc.add_paragraph()
    mod.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = mod.add_run('\nPROGRAMACIÓN DE INTELIGENCIA ARTIFICIAL')
    run.bold = True
    run.size = Pt(18)

    doc.add_paragraph('\n' * 4)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Sistema de Clasificación Automática de Catálogo mediante Visión Artificial\nImplementación con Amazon Rekognition y FastAPI')
    run.size = Pt(14)
    run.italic = True

    doc.add_paragraph('\n' * 8)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info.add_run('Alumno: [Tu Nombre]\nDocente: Equipo Medac\nCurso: Especialización IA y Big Data\nFecha: Abril 2026')

    doc.add_page_break()

    # 2. INTRODUCCIÓN Y OBJETIVOS (Página 2)
    doc.add_heading('1. Introducción y Contextualización', level=1)
    doc.add_paragraph(
        "La empresa FastRetail, en su proceso de expansión digital, ha identificado una ineficiencia crítica "
        "en la gestión de su catálogo online. El etiquetado manual de productos de moda no solo consume una "
        "cantidad excesiva de recursos humanos, sino que introduce errores de categorización que afectan "
        "directamente a la experiencia del usuario y al SEO de la plataforma."
    )
    doc.add_paragraph(
        "Este proyecto presenta una solución basada en Inteligencia Artificial y computación en la nube para "
        "automatizar el proceso de clasificación de imágenes. Utilizando servicios de vanguardia como "
        "Amazon Rekognition, se ha desarrollado una herramienta capaz de analizar visualmente los nuevos "
        "artículos y asignarles etiquetas descriptivas con un alto grado de fiabilidad."
    )

    doc.add_heading('1.1. Objetivos del Proyecto', level=2)
    p = doc.add_paragraph()
    p.add_run("● Desarrollar un sistema de clasificación automática utilizando Visión Artificial.\n")
    p.add_run("● Implementar un pipeline de datos robusto entre el frontend y la infraestructura AWS.\n")
    p.add_run("● Garantizar la escalabilidad mediante el uso de servicios Cloud (S3 y Rekognition).\n")
    p.add_run("● Automatizar la persistencia de datos mediante una base de datos SQL para auditoría y consulta.")

    doc.add_page_break()

    # 3. ANÁLISIS Y DISEÑO DEL MODELO (Página 3)
    doc.add_heading('2. Análisis y Diseño del Modelo de IA', level=1)
    doc.add_paragraph(
        "El núcleo del sistema reside en la integración con Amazon Rekognition. Se ha seleccionado este "
        "modelo por su capacidad de aprendizaje profundo (Deep Learning) pre-entrenado, que ofrece una "
        "precisión superior en la detección de objetos y conceptos de moda sin necesidad de un entrenamiento "
        "costoso desde cero."
    )

    doc.add_heading('2.1. Definición de Requisitos', level=2)
    doc.add_paragraph(
        "Se han definido requisitos técnicos estrictos para asegurar el éxito del sistema:\n"
        "- Umbral de confianza mínimo del 75% para etiquetas automáticas.\n"
        "- Soporte multiformato (incluyendo WebP y PNG) mediante preprocesado local.\n"
        "- Tiempo de respuesta inferior a 2 segundos para el análisis visual.\n"
        "- Almacenamiento persistente del histórico de análisis para mejorar el catálogo futuro."
    )

    doc.add_heading('2.2. Arquitectura de la Solución', level=2)
    doc.add_paragraph(
        "La arquitectura sigue un modelo de microservicio moderno:\n"
        "1. Cliente (Frontend): Interfaz web en HTML5/Vanilla JS para interactuar con el usuario.\n"
        "2. API (Backend): Desarrollada en FastAPI (Python) por su alto rendimiento y tipado asíncrono.\n"
        "3. Procesamiento Local: Uso de la librería Pillow para la normalización de formatos de imagen.\n"
        "4. Servicios Cloud: Amazon S3 para el almacenamiento de archivos y Rekognition para la inferencia de IA.\n"
        "5. Base de Datos: SQLite para la gestión local de metadatos."
    )

    doc.add_page_break()

    # 4. IMPLEMENTACIÓN TÉCNICA (Página 4)
    doc.add_heading('3. Implementación Técnica y Código Fuente', level=1)
    doc.add_paragraph(
        "En esta sección se detalla la implementación del código fuente, destacando la integración con el "
        "SDK Boto3 de Amazon y la gestión de la base de datos mediante SQLAlchemy."
    )

    doc.add_heading('3.1. Lógica del Backend (FastAPI)', level=2)
    doc.add_paragraph("[INSERTAR AQUÍ CAPTURA DEL CÓDIGO DE MAIN.PY - SECCIÓN DE CONFIGURACIÓN Y RUTAS]")
    doc.add_paragraph(
        "El backend gestiona el ciclo de vida de la imagen: recibe el buffer, lo normaliza a JPEG para "
        "garantizar compatibilidad, lo sube a un bucket de S3 y finalmente solicita el análisis a Rekognition."
    )

    doc.add_heading('3.2. Gestión de Datos y Persistencia', level=2)
    doc.add_paragraph("[INSERTAR AQUÍ CAPTURA DEL CÓDIGO DE DATABASE.PY]")
    doc.add_paragraph(
        "Se ha utilizado un ORM (Object-Relational Mapping) para definir el esquema de la tabla 'product_images', "
        "donde se almacenan las etiquetas generadas en formato JSON, permitiendo una fácil recuperación posterior."
    )

    doc.add_paragraph("\nRepositorio de GitHub: https://github.com/[TuUsuario]/fastretail-ai-pfa")

    doc.add_page_break()

    # 5. EJECUCIÓN Y PRUEBAS (Páginas 5-6)
    doc.add_heading('4. Ejecución del Sistema y Validación', level=1)
    doc.add_paragraph(
        "Para validar el sistema, se han realizado pruebas con diferentes categorías de productos (calzado, "
        "accesorios y ropa técnica). Los resultados demuestran que la IA identifica correctamente no solo "
        "el objeto principal, sino también atributos secundarios como color y estilo."
    )
    
    doc.add_heading('4.1. Interfaz de Usuario en Funcionamiento', level=2)
    doc.add_paragraph("[INSERTAR AQUÍ CAPTURA DE LA WEB PRINCIPAL CON VARIAS IMÁGENES CARGADAS]")
    doc.add_paragraph(
        "La interfaz muestra de forma dinámica las etiquetas devueltas por el servicio de Amazon, indicando "
        "visualmente la categorización automática realizada por el sistema."
    )

    doc.add_heading('4.2. Validación en la Consola AWS', level=2)
    doc.add_paragraph("[INSERTAR AQUÍ CAPTURA DEL BUCKET DE S3 DONDE SE VEAN LOS ARCHIVOS SUBIDOS]")
    doc.add_paragraph(
        "Se confirma que el almacenamiento en la nube funciona correctamente, manteniendo una copia de "
        "seguridad de cada producto analizado con nombres de archivo únicos (UUID) para evitar conflictos."
    )

    doc.add_page_break()

    # 6. PROBLEMAS, ÉTICA Y BIBLIOGRAFÍA (Página 7)
    doc.add_heading('5. Problemas Encontrados y Soluciones', level=1)
    doc.add_paragraph(
        "Durante el desarrollo surgieron varios desafíos técnicos:\n"
        "1. Incompatibilidad de Formatos: Rekognition fallaba con archivos WebP. Se solucionó implementando "
        "un conversor automático con la librería Pillow que transforma cualquier entrada a JPEG de alta calidad.\n"
        "2. Gestión de Credenciales en Academia: Al usar AWS Academy, las credenciales son temporales. Se "
        "implementó un sistema de variables de entorno (.env) que permite actualizar el Session Token fácilmente."
    )

    doc.add_heading('6. Consideraciones Éticas y Legales', level=1)
    doc.add_paragraph(
        "El sistema cumple con el RGPD al tratar datos de productos y no personales. Además, se aplican los "
        "principios de transparencia en IA al mostrar al operario humano el nivel de confianza de cada etiqueta, "
        "permitiendo una supervisión informada de la automatización."
    )

    doc.add_heading('7. Bibliografía', level=1)
    doc.add_paragraph(
        "1. AWS. (2026). Amazon Rekognition Developer Guide. https://docs.aws.amazon.com/rekognition/\n"
        "2. Starlette & FastAPI. (2026). TemplateResponse documentation. https://fastapi.tiangolo.com/\n"
        "3. MEDAC. (2026). Contenidos del Módulo de Programación de IA y Big Data."
    )

    # Guardar el documento
    file_path = 'Reporte_FastRetail_AI.docx'
    doc.save(file_path)
    print(f"Reporte extendido generado exitosamente: {file_path}")

if __name__ == "__main__":
    create_report()
